import os
from django.core.management.base import BaseCommand
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

import os
import re
import json
import PyPDF2
import cohere

from django.core.management.base import BaseCommand
from chatbot.models import PolicyFAQ  # Imported per instruction, though not used here.

class Command(BaseCommand):
    help = "Generate a PDF and DOCX file containing all FAQs"

    def handle(self, *args, **kwargs):
        faqs = PolicyFAQ.objects.all()

        # File Paths
        output_dir = "media/faq_documents"
        os.makedirs(output_dir, exist_ok=True)
        pdf_path = os.path.join(output_dir, "FAQs.pdf")
        docx_path = os.path.join(output_dir, "FAQs.docx")

        # Generate PDF
        self.generate_pdf(faqs, pdf_path)

        # Generate DOCX
        self.generate_docx(faqs, docx_path)

        self.stdout.write(self.style.SUCCESS(f"PDF and DOCX generated successfully in {output_dir}"))

    def generate_pdf(self, faqs, file_path):
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        y_position = height - 40

        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y_position, "FAQs Document")
        y_position -= 30

        c.setFont("Helvetica", 12)
        for idx, faq in enumerate(faqs, start=1):
            if y_position < 50:  # Add new page if needed
                c.showPage()
                y_position = height - 40

            c.setFont("Helvetica-Bold", 12)
            c.drawString(40, y_position, f"{idx}. {faq.question}")
            y_position -= 20

            c.setFont("Helvetica", 11)
            for line in faq.answer.split("\n"):
                c.drawString(60, y_position, line)
                y_position -= 15

            y_position -= 10  # Space between questions

        c.save()

    def generate_docx(self, faqs, file_path):
        doc = Document()
        doc.add_heading("FAQs Document", level=1)

        for idx, faq in enumerate(faqs, start=1):
            doc.add_paragraph(f"{idx}. {faq.question}", style="Heading 2")
            doc.add_paragraph(faq.answer, style="Normal")
            doc.add_paragraph("")  # Space between questions

        doc.save(file_path)
