from docx import Document
from docx.shared import Pt

# Create a new Document.
document = Document()

# Add a title.
document.add_heading("Architecture & Implementation Document for a Complex Query Answering System", 0)

# Section 1: Overview
document.add_heading("1. Overview", level=1)
document.add_paragraph(
    "Objective:\n"
    "Build a system that handles complex, non-repetitive queries by students. Instead of simply returning long excerpts, "
    "the system will retrieve relevant, detailed information from a corpus and generate a concise, easy-to-understand answer."
)
document.add_paragraph(
    "Approach:\n"
    "Use a Retrieval-Augmented Generation (RAG) pipeline that:\n"
    "- Retrieves relevant document sections from a large corpus.\n"
    "- Generates a simplified and summarized answer using a fine-tuned generative model."
)

# Section 2: Data Storage and Format
document.add_heading("2. Data Storage and Format", level=1)
document.add_heading("Document Corpus (Detailed Information):", level=2)
document.add_paragraph(
    "Format: Store your data as detailed documents or paragraphs.\n"
    "Structure: Each document should include:\n"
    "  - Title: A short, descriptive title.\n"
    "  - Content: The full, detailed explanation.\n"
    "  - Metadata (Optional): Tags, topics, or section headings that aid in retrieval."
)
document.add_paragraph("Example JSON Format:")
document.add_paragraph(
    """[
  {
    "title": "Advanced Quantum Mechanics",
    "content": "Quantum mechanics is a branch of physics that explains the behavior of matter and energy at the atomic and subatomic levels. It involves complex concepts such as wave-particle duality, uncertainty principles, and quantum entanglement...",
    "metadata": {"tags": ["physics", "quantum", "advanced"]}
  },
  {
    "title": "Effective Study Techniques for Complex Subjects",
    "content": "When dealing with complex topics, breaking down information into smaller, manageable segments is crucial. Techniques such as spaced repetition, summarization, and active recall can significantly improve understanding...",
    "metadata": {"tags": ["education", "study techniques", "complex subjects"]}
  }
]""",
    style="Intense Quote"
)
document.add_paragraph(
    "Training Data for Simplification (Optional but Recommended):\n"
    "Pair each complex document (or excerpt) with a concise summary."
)
document.add_paragraph("Example Pair:")
document.add_paragraph(
    """{
  "input": "Quantum mechanics is a branch of physics that explains the behavior of matter and energy at the atomic and subatomic levels, involving complex principles like wave-particle duality and uncertainty.",
  "output": "Quantum mechanics studies how tiny particles behave in ways that differ from everyday objects, with principles that challenge common intuition."
}""",
    style="Intense Quote"
)

# Section 3: System Modules
document.add_heading("3. System Modules", level=1)
document.add_heading("A. Data Ingestion & Preprocessing:", level=2)
document.add_paragraph(
    "Storage: Use JSON files or a database (via Django models) to store your documents.\n"
    "Preprocessing: Use spaCy to tokenize and clean the text, and Sentence Transformers to encode the text.\n"
    "Indexing: Build a FAISS index for vector similarity search; optionally, use BM25 for keyword-based retrieval."
)
document.add_heading("B. Retrieval Module:", level=2)
document.add_paragraph(
    "Function: Accept a student's query and retrieve relevant document sections using BM25 and/or FAISS.\n"
    "Outcome: Provide rich context to the generative model."
)
document.add_heading("C. Generative Model (Simplification & Answering):", level=2)
document.add_paragraph(
    "Model Options: Pre-trained models like T5, FLAN-T5, or GPT-based models.\n"
    "Fine-Tuning: Fine-tune on your paired data (complex input to simplified output) if available.\n"
    "Prompt Engineering: Use prompts that instruct the model to provide a clear, concise answer."
)
document.add_heading("D. Chatbot Interface and Integration:", level=2)
document.add_paragraph(
    "Telegram Bot: Use the Telegram Bot API (with python-telegram-bot) and store sensitive information (e.g., bot token) in Django settings.\n"
    "Django Backend: Manage data and configuration.\n"
    "(Optional) Rasa Integration: For additional NLU and conversation management."
)

# Section 4: Tools & Technologies
document.add_heading("4. Tools & Technologies", level=1)
document.add_paragraph(
    "- Programming Language: Python (3.8 or 3.9 recommended)\n"
    "- Frameworks: Django for backend; Telegram Bot API and optionally Rasa for the chatbot interface.\n"
    "- NLP Libraries: spaCy, Sentence Transformers, Transformers (Hugging Face)\n"
    "- Indexing: FAISS for vector search; BM25 for keyword retrieval.\n"
    "- Environment: python-dotenv for managing environment variables."
)

# Section 5: Implementation Steps
document.add_heading("5. Implementation Steps", level=1)
document.add_heading("Step 1: Environment Setup", level=2)
document.add_paragraph(
    "Create and activate a virtual environment.\n"
    "Install dependencies:\n"
    "    pip install django rasa transformers sentencepiece spacy faiss-cpu rank_bm25 python-dotenv\n"
    "    python -m spacy download en_core_web_sm"
)
document.add_heading("Step 2: Data Ingestion", level=2)
document.add_paragraph("Store and load your detailed documents (JSON or Django models).")
document.add_heading("Step 3: Indexing and Retrieval", level=2)
document.add_paragraph(
    "Generate embeddings using Sentence Transformers.\n"
    "Build a FAISS index and optionally a BM25 index."
)
document.add_heading("Step 4: Generative Module", level=2)
document.add_paragraph(
    "Load and optionally fine-tune a pre-trained model (e.g., FLAN-T5).\n"
    "Craft prompts for simplified answer generation."
)
document.add_heading("Step 5: Integration", level=2)
document.add_paragraph(
    "Integrate with a Telegram bot (import token from Django settings).\n"
    "Combine retrieval and generation to provide answers."
)
document.add_heading("Step 6: Deployment and Testing", level=2)
document.add_paragraph(
    "Run the Django server.\n"
    "Start the Telegram bot.\n"
    "Test the system with sample complex queries."
)

# Section 6: System Diagram (Conceptual)
document.add_heading("6. System Diagram (Conceptual)", level=1)
document.add_paragraph("         ┌─────────────────┐")
document.add_paragraph("         │  Student Query  │")
document.add_paragraph("         └───────┬─────────┘")
document.add_paragraph("                 │")
document.add_paragraph("        ┌────────▼─────────┐")
document.add_paragraph("        │ Retrieval Module │  <-- Uses BM25 & FAISS")
document.add_paragraph("        └────────┬─────────┘")
document.add_paragraph("                 │")
document.add_paragraph("        ┌────────▼─────────┐")
document.add_paragraph("        │ Generative Model │  <-- Simplifies & summarizes answer")
document.add_paragraph("        └────────┬─────────┘")
document.add_paragraph("                 │")
document.add_paragraph("         ┌───────▼─────────┐")
document.add_paragraph("         │   Answer Output  │")
document.add_paragraph("         └─────────────────┘")

# Section 7: Conclusion
document.add_heading("7. Conclusion", level=1)
document.add_paragraph(
    "By employing a hybrid Retrieval-Augmented Generation approach:\n"
    "- Complex information is preserved from detailed documents.\n"
    "- Concise, simplified answers are generated for clarity.\n"
    "- Modern NLP techniques and a modular design ensure scalability and accuracy.\n\n"
    "This design will help answer challenging, non-repetitive student queries in a clear, understandable manner."
)

# Save the document.
document.save("Complex_Query_Answering_System.docx")
print("Word document generated successfully as 'Complex_Query_Answering_System.docx'.")
